import logging
from logging.handlers import RotatingFileHandler
import os
import importlib
from flask import Flask, render_template, request, jsonify, flash, redirect, url_for, session, current_app
import json
from datetime import datetime, timedelta
import uuid
import sqlite3
import bleach
from werkzeug.security import generate_password_hash, check_password_hash
import secrets
from functools import wraps
import PyPDF2
import mimetypes
from werkzeug.utils import secure_filename
import requests
from utils.weather_service import WeatherService, WaterOutageService
from utils.regional_service import RegionalService
from utils.aws_bill_analyzer import AWSBillAnalyzer, allowed_file
from utils.advanced_bill_analyzer import process_bill_with_advanced_analysis
from utils.intelligent_bill_analyzer import analyze_bill
from utils.water_tax_expert import create_water_expert
from utils.water_bill_chat import create_water_bill_chat
import traceback
from weather_routes import weather_bp
import boto3
from botocore.exceptions import ClientError, NoCredentialsError, CredentialRetrievalError
from ollama_data_analysis import OllamaDataAnalyzer

# Configure logging
log_dir = os.path.join(os.path.dirname(__file__), 'logs')
os.makedirs(log_dir, exist_ok=True)
log_file = os.path.join(log_dir, 'app.log')

# Create a global logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Create file and console handlers
file_handler = logging.FileHandler(log_file)
console_handler = logging.StreamHandler()

# Create formatter
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

# Add handlers to logger
logger.addHandler(file_handler)
logger.addHandler(console_handler)

app = Flask(__name__, 
            template_folder='templates', 
            static_folder='static')
app.logger.addHandler(file_handler)
app.logger.addHandler(console_handler)
app.logger.setLevel(logging.INFO)
app.secret_key = os.urandom(24)  
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  
app.config['SESSION_COOKIE_SECURE'] = False  
app.config['SESSION_COOKIE_HTTPONLY'] = True  
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  
app.config['SERVER_NAME'] = None  

# Uploads directory configuration
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'csv'}

# Initialize weather, outage and regional services
weather_service = WeatherService()
outage_service = WaterOutageService()
regional_service = RegionalService()

# Configure AWS Bill Analyzer
aws_bill_analyzer = AWSBillAnalyzer(app.config['UPLOAD_FOLDER'])  

def initialize_chatbot():
    """
    Global chatbot initialization function with simplified, robust model loading
    """
    import importlib.util
    import sys
    import requests
    import traceback

    current_dir = os.path.dirname(os.path.abspath(__file__))
    sys.path.insert(0, current_dir)
    sys.path.insert(0, os.path.join(current_dir, 'utils'))
    
    # Simplified Ollama API URL
    ollama_api_url = os.getenv('OLLAMA_API_URL', 'http://localhost:11434')
    
    def check_ollama_model(model_name, max_attempts=3):
        """Simplified model availability check with retry mechanism"""
        for attempt in range(max_attempts):
            try:
                # Quick timeout for model check
                response = requests.get(f'{ollama_api_url}/api/tags', timeout=10)
                
                if response.status_code == 200:
                    # Minimal logging
                    app.logger.info(f"Checking model: {model_name}")
                    
                    # Simple model existence check
                    models = response.json().get('models', [])
                    if any(model_name in model.get('name', '') for model in models):
                        return True
                
                # Attempt to pull model if not found
                pull_response = requests.post(
                    f'{ollama_api_url}/api/pull', 
                    json={'name': model_name, 'stream': False},
                    timeout=60
                )
                
                if pull_response.status_code == 200:
                    app.logger.info(f"Successfully pulled model: {model_name}")
                    return True
            
            except Exception as e:
                app.logger.warning(f"Model check attempt {attempt + 1} failed: {e}")
                time.sleep(2 ** attempt)  # Exponential backoff
        
        return False

    # Simplified model check
    check_ollama_model('llama3.2')
    
    # Import strategies with minimal logging
    import_strategies = [
        ('model_inference', 'WaterConservationBot'),
        ('utils.water_tax_expert', 'WaterTaxExpert')
    ]
    
    for module_name, class_name in import_strategies:
        try:
            module_path = os.path.join(current_dir, f"{module_name.replace('.', '/')}.py")
            
            if os.path.exists(module_path):
                spec = importlib.util.spec_from_file_location(module_name, module_path)
                module = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(module)
                
                if hasattr(module, class_name):
                    chatbot_class = getattr(module, class_name)
                    return chatbot_class()
        
        except Exception as e:
            app.logger.error(f"Import error for {module_name}.{class_name}: {e}")
    
    # Fallback mechanism
    app.logger.critical("Could not initialize any chatbot")
    return None

# Initialize chatbot during app startup
with app.app_context():
    app.chatbot = initialize_chatbot()

# Database initialization
def init_db():
    conn = sqlite3.connect('feedback.db')
    c = conn.cursor()
    
    # Create feedback table
    c.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            email TEXT,
            rating INTEGER,
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create admin table
    c.execute('''
        CREATE TABLE IF NOT EXISTS admin (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT
        )
    ''')
    
    # Add default admin if not exists
    c.execute('SELECT * FROM admin WHERE username = ?', ('admin',))
    if not c.fetchone():
        c.execute('INSERT INTO admin (username, password) VALUES (?, ?)',
                 ('admin', generate_password_hash('admin123')))
    
    conn.commit()
    conn.close()

# Initialize database on startup
init_db()

# Load translations
with open('translations.json', 'r', encoding='utf-8') as f:
    translations = json.load(f)

def get_text(key):
    """Get text in English"""
    try:
        return translations['en'].get(key, key)
    except:
        return key

# Register get_text as a template filter
app.jinja_env.filters['translate'] = get_text

@app.context_processor
def utility_processor():
    return dict(get_text=get_text)

# Sanitize input
def sanitize_input(text):
    return bleach.clean(text, strip=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_type(filename):
    """Get file type based on extension"""
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    return mimetypes.guess_type(filename)[0] or 'application/octet-stream'

def read_file_content(file_path):
    """
    Read file content based on file type with robust error handling
    Supports multiple file types: txt, pdf, docx, images
    
    Args:
        file_path (str): Path to the uploaded file
    
    Returns:
        str: Extracted text content from the file
    """
    try:
        # Determine file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        # Text files
        if ext in ['.txt', '.csv']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        # PDF files
        elif ext == '.pdf':
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text

        # Word documents
        elif ext in ['.doc', '.docx']:
            import docx
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs if para.text])

        # Image files (OCR)
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            import pytesseract
            from PIL import Image
            
            # Ensure Tesseract is installed
            try:
                return pytesseract.image_to_string(Image.open(file_path))
            except Exception as ocr_error:
                logger.warning(f"OCR failed for image: {ocr_error}")
                return "Could not extract text from image"

        else:
            logger.warning(f"Unsupported file type: {ext}")
            return f"Unsupported file type: {ext}"

    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        return f"Error reading file: {str(e)}"

def generate_csrf_token():
    if 'csrf_token' not in session:
        session['csrf_token'] = secrets.token_hex(32)
    return session['csrf_token']

app.jinja_env.globals['csrf_token'] = generate_csrf_token

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in to access this page', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def verify_csrf_token():
    # Skip CSRF check for local network requests
    if request.remote_addr.startswith('192.168.') or request.remote_addr == '127.0.0.1':
        return True
        
    # Check form data first
    token = request.form.get('csrf_token')
    if not token and request.is_json:
        # Then check JSON data
        token = request.json.get('csrf_token')
    # Finally check headers
    if not token:
        token = request.headers.get('X-CSRF-Token')
    return token and token == session.get('csrf_token')

@app.before_request
def before_request():
    # Skip session check for static files and public routes
    if request.path.startswith('/static') or request.endpoint in ['login', 'index', 'logout']:
        return

    # Only check admin routes
    if request.endpoint and request.endpoint.startswith('admin'):
        if 'user_id' not in session:
            flash('Please log in to access the admin panel.', 'error')
            return redirect(url_for('login'))
        
        # Check session timeout (20 minutes)
        if 'last_activity' in session:
            last_activity = datetime.fromisoformat(session['last_activity'])
            if (datetime.now() - last_activity).total_seconds() > 1200:  # 20 minutes
                session.clear()
                flash('Your session has expired. Please log in again.', 'error')
                return redirect(url_for('login'))
        
        # Update last activity time
        session['last_activity'] = datetime.now().isoformat()

    # Require CSRF token for all POST requests
    if request.method == 'POST':
        if not verify_csrf_token():
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'error': 'Invalid CSRF token'}), 400
            flash('Invalid CSRF token. Please try again.', 'error')
            return redirect(url_for('index'))

@app.route('/')
def index():
    try:
        logger.debug("Rendering index.html")
        return render_template('index.html')
    except Exception as e:
        logger.error(f"Error rendering index.html: {str(e)}")
        return str(e), 500

@app.route('/login', methods=['GET', 'POST'])
def login():
    # If already logged in, redirect to admin
    if 'user_id' in session:
        return redirect(url_for('admin'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if username == 'admin' and password == 'admin123':
            session.clear()  # Clear any existing session
            session['user_id'] = username
            session['last_activity'] = datetime.now().isoformat()
            flash('Welcome back!', 'success')
            return redirect(url_for('admin'))
        else:
            flash('Invalid username or password', 'error')
            return render_template('login.html', error='Invalid username or password')
    
    return render_template('login.html')

@app.route('/admin')
@login_required
def admin():
    try:
        if 'user_id' not in session:
            flash('Please log in to access the admin panel.', 'error')
            return redirect(url_for('login'))

        conn = sqlite3.connect('feedback.db')
        c = conn.cursor()
        
        # Fetch feedback data
        feedbacks = c.execute('''
            SELECT id, name, email, rating, comments, timestamp 
            FROM feedback 
            ORDER BY timestamp DESC
        ''').fetchall()
        
        # Convert to list of dictionaries
        feedback_list = []
        for f in feedbacks:
            feedback_list.append({
                'id': f[0],
                'name': f[1],
                'email': f[2],
                'rating': f[3],
                'comments': f[4],
                'timestamp': f[5]
            })
        
        conn.close()
        return render_template('admin.html', feedbacks=feedback_list)
        
    except Exception as e:
        logger.error(f"Error in admin endpoint: {str(e)}")
        flash('An error occurred while loading the admin panel', 'error')
        return redirect(url_for('index'))

@app.route('/admin/delete_feedback/<int:feedback_id>', methods=['POST'])
@login_required
def admin_delete_feedback(feedback_id):
    try:
        # Get request data
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No JSON data provided'}), 400

        # Verify CSRF token from request body
        if not data.get('csrf_token') or data['csrf_token'] != session.get('csrf_token'):
            return jsonify({'success': False, 'error': 'Invalid CSRF token'}), 400
            
        conn = sqlite3.connect('feedback.db')
        c = conn.cursor()
        
        # Check if feedback exists
        feedback = c.execute('SELECT id FROM feedback WHERE id = ?', (feedback_id,)).fetchone()
        
        if feedback:
            c.execute('DELETE FROM feedback WHERE id = ?', (feedback_id,))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Feedback deleted successfully'})
        else:
            conn.close()
            return jsonify({'success': False, 'error': 'Feedback not found'}), 404
            
    except sqlite3.Error as e:
        if 'conn' in locals():
            conn.close()
        return jsonify({'success': False, 'error': f'Database error: {str(e)}'}), 500
    except Exception as e:
        if 'conn' in locals():
            conn.close()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/feedback_structure')
@login_required
def get_feedback_structure():
    try:
        conn = sqlite3.connect('feedback.db')
        c = conn.cursor()
        
        # Get table info
        table_info = c.execute("PRAGMA table_info(feedback)").fetchall()
        
        # Get a sample feedback
        sample = c.execute("SELECT * FROM feedback LIMIT 1").fetchone()
        
        conn.close()
        
        return jsonify({
            'table_structure': [{'name': col[1], 'type': col[2]} for col in table_info],
            'sample_data': sample
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out successfully.', 'success')
    return redirect(url_for('index'))

@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['rating', 'name', 'email', 'comment']
        for field in required_fields:
            if field not in data:
                return jsonify({'status': 'error', 'message': f'Missing required field: {field}'}), 400
        
        # Sanitize inputs
        name = sanitize_input(data['name'])
        email = sanitize_input(data['email'])
        comment = sanitize_input(data['comment'])
        rating = int(data['rating'])
        
        # Validate rating
        if not 1 <= rating <= 5:
            return jsonify({'status': 'error', 'message': 'Rating must be between 1 and 5'}), 400
        
        conn = sqlite3.connect('feedback.db')
        c = conn.cursor()
        c.execute('''
            INSERT INTO feedback (name, email, rating, comments)
            VALUES (?, ?, ?, ?)
        ''', (name, email, rating, comment))
        conn.commit()
        conn.close()
        
        return jsonify({'status': 'success'})
    except Exception as e:
        logger.error(f"Error submitting feedback: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/update_status', methods=['POST'])
@login_required
def update_status():
    try:
        if not verify_csrf_token():
            return jsonify({'status': 'error', 'message': 'Invalid CSRF token'}), 400
        
        feedback_id = request.form.get('feedback_id')
        new_status = request.form.get('status')
        
        if not feedback_id or not new_status:
            return jsonify({'status': 'error', 'message': 'Missing required fields'}), 400
        
        conn = sqlite3.connect('feedback.db')
        c = conn.cursor()
        c.execute('UPDATE feedback SET status = ? WHERE id = ?', (new_status, feedback_id))
        conn.commit()
        conn.close()
        
        return jsonify({'status': 'success', 'message': f'Status updated to {new_status}'})
    except Exception as e:
        logger.error(f"Error in update_status endpoint: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/export_to_training', methods=['POST'])
def export_to_training():
    try:
        if not verify_csrf_token():
            return jsonify({'status': 'error', 'message': 'Invalid CSRF token'})
        
        feedback_id = request.form.get('feedback_id')
        
        # Get feedback from database
        conn = sqlite3.connect('feedback.db')
        c = conn.cursor()
        c.execute('SELECT * FROM feedback WHERE id = ?', (feedback_id,))
        feedback = c.fetchone()
        conn.close()
        
        if not feedback:
            return jsonify({'status': 'error', 'message': 'Feedback not found'})
        
        # Format feedback for training
        training_data = {
            "messages": [
                {"role": "user", "content": feedback[4]},  
                {"role": "assistant", "content": ""}  
            ]
        }
        
        # Append to training dataset
        training_file = os.path.join('dataset', 'training_feedback.jsonl')
        with open(training_file, 'a', encoding='utf-8') as f:
            json.dump(training_data, f, ensure_ascii=False)
            f.write('\n')
        
        # Update feedback status
        conn = sqlite3.connect('feedback.db')
        c = conn.cursor()
        c.execute('UPDATE feedback SET status = ? WHERE id = ?', ('exported', feedback_id))
        conn.commit()
        conn.close()
        
        return jsonify({'status': 'success'})
    except Exception as e:
        logger.error(f"Error in export_to_training endpoint: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        # Log all incoming files
        logger.info(f"Received files: {request.files}")
        
        if 'file' not in request.files:
            logger.warning('No file part in the request')
            return jsonify({'response': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        # Log file details
        logger.info(f"File name: {file.filename}")
        logger.info(f"File content type: {file.content_type}")
        
        if file.filename == '':
            logger.warning('No selected file')
            return jsonify({'response': 'No file selected'}), 400
        
        # Detaylı dosya türü kontrolü
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        logger.info(f"File extension: {file_ext}")
        
        if file and file_ext in ALLOWED_EXTENSIONS:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # Ensure upload directory exists
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            
            # Save file
            file.save(filepath)
            logger.info(f"File saved to: {filepath}")
            
            # Read file content with detailed logging
            try:
                content = read_file_content(filepath)
                logger.info(f"File content length: {len(content)} characters")
                
                if not content:
                    logger.warning(f"No content extracted from {filename}")
                    return jsonify({'response': 'No content extracted'}), 400
            except Exception as read_error:
                logger.error(f"File reading error: {read_error}")
                return jsonify({'response': 'No content extracted'}), 500
            
            # Clean up the uploaded file
            os.remove(filepath)
            logger.info(f"Temporary file {filepath} removed")
            
            # Generate response about the document content
            prompt = f"I will provide you with a document content. Please analyze it and identify any water conservation related information, tips, or relevant content. If there are water conservation practices mentioned, summarize them. If there's no water-related content, suggest how the topic could be connected to water conservation. Here's the content:\n\n{content}"
            
            try:
                response, error = app.chatbot.generate_response(prompt)
                if error:
                    logger.error(f"Chatbot response error: {error}")
                    response = error
                logger.info(f"Generated response length: {len(response) if response else 0}")
            except Exception as chat_error:
                logger.error(f"Chatbot error: {str(chat_error)}")
                response = "Sorry, I encountered an error analyzing the document."
            
            return jsonify({'response': response})
        
        logger.warning(f'Invalid file type: {file_ext}')
        return jsonify({'response': 'Invalid file type. Please upload a .txt, .pdf, .csv file.'}), 400
    
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        return jsonify({'response': 'Sorry, I encountered an error analyzing the document. Please try again.'})

def chat_api(role):
    """
    Enhanced chat endpoint with specialized role-based responses
    
    Supports:
    - General conversation
    - Specialized role-based interactions
    - Bill analysis and explanation
    - File upload handling
    """
    try:
        # Dynamic model selection based on context
        from model_inference_education import WaterConservationBot as EducationBot
        from model_inference_farmers import WaterConservationBot as FarmersBot
        from model_inference import WaterConservationBot as WaterBot
        # Determine the appropriate bot based on the role
        if role == 'farmer':
            water_expert = FarmersBot(model_name='water-expert-farmers')
        elif role == 'educator':
            water_expert = EducationBot(model_name='water-expert-education')
        elif role == "expert":
            water_expert = WaterBot(model_name='water-expert')
        
        # Handle both JSON and form data
        message = None
        file_content = None
        
        # Check if it's a file upload or JSON request
        if request.is_json:
            data = request.get_json()
            message = data.get('message', '').strip()
        else:
            # Form data (for file uploads)
            message = request.form.get('message', '').strip()
            
            # Check for file upload
            if 'file' in request.files:
                uploaded_file = request.files['file']
                if uploaded_file and allowed_file(uploaded_file.filename):
                    # Save file temporarily
                    filename = secure_filename(uploaded_file.filename)
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    
                    # Save file
                    uploaded_file.save(file_path)

                    # Read file content
                    file_content = read_file_content(file_path)
                    
                    # Optional: Remove the temporary file
                    os.remove(file_path)
        
        if not message:
            return jsonify({"error": "No message provided"}), 400
        
        # Combine message with file content if available
        if file_content:
            message = f"{message}\n\nFile Content:\n{file_content}"
        
        # Generate response using the appropriate bot
        response, error = water_expert.generate_response(message)
        
        if error:
            return jsonify({"error": str(error)}), 500
        
        return jsonify({"response": response})
    
    except Exception as e:
        logger.error(f"Error processing message: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            'error': str(e),
            'response': 'An unexpected error occurred.'
        }), 500

@app.route('/api/main-chat', methods=['POST'])
def main_chat_api():
    return chat_api('main')

@app.route('/api/farmer-chat', methods=['POST'])
def farmer_chat_api():
    """
    Specialized chat endpoint for farmer interactions.
    
    Handles:
    - Text messages
    - File uploads
    - Farmer-specific context
    """
    # Log the request details
    logger.info(f"Received request headers: {request.headers}")
    logger.info(f"Request content type: {request.content_type}")
    
    # Check if it's a multipart form data (file upload)
    if 'file' in request.files:
        file = request.files['file']
        logger.info(f"File uploaded: {file.filename}")
        # You can add file processing logic here if needed
    
    # Extract message
    message = request.form.get('message', '').strip()
    logger.info(f"Received message: {message}")
    
    # Validate message
    if not message:
        logger.warning("Empty message received")
        return jsonify({"error": "Message cannot be empty"}), 400
    
    # Ensure farmer-specific chatbot is initialized
    if not hasattr(app, 'farmer_chatbot'):
        from model_inference_farmers import WaterConservationBot
        app.farmer_chatbot = WaterConservationBot()
    
    # Generate response
    try:
        # Verify generate_response method exists
        if not hasattr(app.farmer_chatbot, 'generate_response'):
            logger.error("Chatbot does not have generate_response method")
            return jsonify({"error": "Chatbot method not found"}), 500

        # Detailed logging before method call
        logger.info(f"Calling generate_response with context: {message[:200]}...")
        
        response, error = app.farmer_chatbot.generate_response(message)
        
        # Log generation results
        logger.info(f"Response generated: {response is not None}")
        logger.info(f"Error: {error}")

        # Handle response
        if error:
            logger.warning(f"Generation error: {error}")
            return jsonify({"error": error}), 400
        
        if not response:
            logger.warning("Empty response generated")
            response = "Üzgünüm, bu soruya şu anda yanıt veremiyorum."
        
        return jsonify({
            "message": response,
            "source": getattr(app.farmer_chatbot, 'model_name', 'water-expert-farmers')
        })
    
    except Exception as e:
        # Log the error (consider using proper logging in production)
        logger.error(f"Unexpected error in farmer chat: {e}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return jsonify({"error": "Internal server error during response generation"}), 500

@app.route('/api/educator-chat', methods=['POST'])
def educator_chat_api():
    """
    Render the educator chat interface and handle chat interactions
    """
    try:
        # Dynamic model selection based on context
        from model_inference_education import WaterConservationBot as EducationBot
        from model_inference_farmers import WaterConservationBot as FarmersBot
        
        # Determine the appropriate bot based on the context
        context = request.json.get('context', 'education')
        
        if context == 'farmers':
            water_expert = FarmersBot(model_name='water-expert-farmers')
        else:
            water_expert = EducationBot(model_name='water-expert-education')
        
        if request.method == 'POST':
            # Handle both JSON and form data
            message = None
            try:
                if request.is_json:
                    data = request.get_json()
                    message = data.get('message', '').strip()
                else:
                    message = request.form.get('message', '').strip()
                
                if not message:
                    return jsonify({"error": "No message provided"}), 400
                
                # Generate response using the appropriate bot
                response, error = water_expert.generate_response(message)
                
                if error:
                    return jsonify({"error": str(error)}), 500
                
                return jsonify({"response": response})
            
            except Exception as e:
                logger.error(f"Error processing message: {str(e)}")
                logger.error(f"Traceback: {traceback.format_exc()}")
                return jsonify({
                    "error": "Internal server error", 
                    "details": str(e)
                }), 500
        
    except Exception as e:
        logger.error(f"Unexpected error in educator_chat_api: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            "error": "Unexpected server error", 
            "details": str(e)
        }), 500

@app.route('/chat_page')
def chat_page():
    try:
        if 'messages' in session:
            session.pop('messages')
        
        logger.debug("Rendering main.html")
        return render_template('main.html')
    except Exception as e:
        logger.error(f"Error rendering main.html: {str(e)}")
        return str(e), 500

@app.route('/clear_chat', methods=['POST'])
def clear_chat():
    """Clear chat history when user leaves the page."""
    try:
        # Clear conversation from session
        if 'conversation' in session:
            session.pop('conversation')
        
        # Clear bot instance from session
        if 'bot' in session:
            session.pop('bot')
            
        session.modified = True
        return jsonify({'status': 'success'})
    except Exception as e:
        logger.error(f"Error clearing chat history: {str(e)}")
        return jsonify({'error': 'Failed to clear chat history'}), 500

@app.route('/qr-scanner')
def qr_scanner():
    return render_template('qr_scanner.html')

@app.route('/api/weather', methods=['GET'])
def get_weather():
    """
    Retrieve weather information for the user's location.
    
    Returns:
    - Weather data including current conditions and forecast
    - Handles various error scenarios
    """
    try:
        # Log the start of weather retrieval
        app.logger.info("Starting weather retrieval process")
        
        # Attempt to get user's IP-based location
        try:
            # Use requests to get IP-based location
            location_response = requests.get('https://ipapi.co/json/', timeout=5)
            location_data = location_response.json()
            
            # Log the location data
            app.logger.info(f"IP Location Data: {location_data}")
            
            # Extract latitude and longitude
            latitude = location_data.get('latitude')
            longitude = location_data.get('longitude')
            location_name = f"{location_data.get('city', 'Unknown')}, {location_data.get('country_name', 'Unknown')}"
            
        except Exception as location_error:
            # Log location retrieval error
            app.logger.error(f"Location retrieval error: {location_error}")
            
            # Fallback to default coordinates (e.g., center of a country)
            latitude = 0.0
            longitude = 0.0
            location_name = 'Global Default'
        
        # Log the coordinates being used
        app.logger.info(f"Using coordinates - Latitude: {latitude}, Longitude: {longitude}")
        
        # Construct Open-Meteo API URL
        weather_url = f"https://api.open-meteo.com/v1/forecast?latitude={latitude}&longitude={longitude}&current_weather=true&hourly=temperature_2m,precipitation_probability&daily=temperature_2m_max,temperature_2m_min&timezone=auto"
        
        # Log the constructed URL
        app.logger.info(f"Weather API URL: {weather_url}")
        
        # Fetch weather data
        weather_response = requests.get(weather_url, timeout=10)
        
        # Log the weather response
        app.logger.info(f"Weather API Response Status: {weather_response.status_code}")
        
        # Check for successful response
        if weather_response.status_code != 200:
            app.logger.error(f"Weather API Error: {weather_response.status_code} - {weather_response.text}")
            return {
                'error': 'Unable to fetch weather data',
                'details': f'API returned status {weather_response.status_code}'
            }
        
        # Parse weather data
        weather_data = weather_response.json()
        
        # Log the received weather data
        app.logger.info(f"Received Weather Data: {weather_data}")
        
        # Extract current weather details
        current_weather = weather_data.get('current_weather', {})
        temperature = current_weather.get('temperature', 'N/A')
        windspeed = current_weather.get('windspeed', 'N/A')
        
        # Return comprehensive weather information
        return {
            'success': True,
            'latitude': latitude,
            'longitude': longitude,
            'location': location_name,
            'temperature': temperature,
            'windspeed': windspeed
        }
    
    except Exception as e:
        # Log any unexpected errors
        app.logger.error(f"Unexpected weather retrieval error: {e}")
        app.logger.error(traceback.format_exc())
        
        return {
            'error': 'Failed to retrieve weather information',
            'details': str(e)
        }

@app.route('/api/water-outages')
def get_water_outages():
    """
    Retrieve water outage information for the user's location.
    
    Returns:
    - Water outage data
    - Handles various error scenarios
    """
    try:
        # Get location parameters
        latitude = request.args.get('lat')
        longitude = request.args.get('lon')
        
        # Validate location parameters
        if not latitude or not longitude:
            # Fallback to default location (e.g., Ankara)
            latitude, longitude = 39.9334, 32.8597
        
        try:
            latitude = float(latitude)
            longitude = float(longitude)
        except ValueError:
            logger.warning(f"Invalid coordinates for water outages: lat={latitude}, lon={longitude}")
            return jsonify({
                "error": "Invalid location coordinates",
                "details": "Latitude and longitude must be numeric values"
            }), 400
        
        # Use water outage service
        try:
            water_outage_service = WaterOutageService()
            outages = water_outage_service.get_water_outages(latitude, longitude)
            
            return jsonify({
                "water_outages": outages,
                "location": get_location_details(latitude, longitude)
            })
        
        except Exception as service_error:
            logger.error(f"Water outage service error: {service_error}", exc_info=True)
            return jsonify({
                "error": "Could not fetch water outage information",
                "details": str(service_error)
            }), 500
    
    except Exception as e:
        logger.error(f"Unexpected error in water outages endpoint: {e}", exc_info=True)
        return jsonify({
            "error": "Internal server error",
            "details": str(e)
        }), 500

@app.route('/api/regional-insights')
def get_regional_insights():
    """Get regional insights including weather and dam status"""
    try:
        # Get location from query parameters or default to a central location
        lat = request.args.get('lat', '39.9334')  # Default: Ankara
        lon = request.args.get('lon', '32.8597')
        region_id = request.args.get('region', 'ankara')
        
        insights = regional_service.get_regional_insights(lat, lon, region_id)
        return jsonify(insights)
    except Exception as e:
        return jsonify({
            "error": str(e),
            "message": "Failed to fetch regional insights"
        }), 500

@app.route('/educators')
def educators():
    """Render the educators page"""
    return render_template('educators.html')

@app.route('/farmers')
def farmers():
    """Render the farmers page"""
    return render_template('farmers.html')

@app.route('/farmer-chat', methods=['GET', 'POST'])
def farmer_chat():
    """Handle farmer chat page rendering and message processing"""
    if request.method == 'POST':
        # Process incoming chat message
        try:
            # Get JSON data from the request
            data = request.get_json()
            
            # Extract message
            message = data.get('message', '').strip()
            
            # Validate message
            if not message:
                return jsonify({'error': 'No message provided'}), 400
            
            # Use model inference for farmers
            from model_inference_farmers import generate_farmer_response
            
            # Generate response
            response_text = generate_farmer_response(message)
            
            # Return JSON response
            return jsonify({
                'response': response_text
            })
        
        except Exception as e:
            # Log the error (consider using proper logging in production)
            print(f"Error processing farmer chat message: {e}")
            return jsonify({
                'error': 'An error occurred while processing your message',
                'details': str(e)
            }), 500
    
    # If it's a GET request, render the template
    return render_template('farmer_chat.html')

@app.route('/educator-chat')
def educator_chat():
    return render_template('educator_chat.html')

@app.route('/water-tax')
def water_tax_page():
    """
    Render the water tax reading page
    """
    try:
        logger.debug("Rendering water_tax.html")
        return render_template('water_tax.html')
    except Exception as e:
        logger.error(f"Error rendering water_tax.html: {str(e)}")
        return str(e), 500

@app.route('/api/water-tax', methods=['POST'])
def water_tax_api():
    """
    API endpoint for water tax bill analysis
    Supports file upload and intelligent bill processing
    """
    try:
        # Check if files are present in the request
        if 'water_bill' not in request.files:
            logger.error("No bill uploaded")
            return jsonify({
                'success': False, 
                'error': 'No bill uploaded',
                'results': []
            }), 400

        # Get uploaded files
        uploaded_files = request.files.getlist('water_bill')
        
        # Ensure uploads directory exists (redundant, but good practice)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

        # Process each uploaded file
        results = []
        for uploaded_file in uploaded_files:
            if uploaded_file.filename == '':
                continue  # Skip empty filenames

            # Secure filename and save
            filename = secure_filename(uploaded_file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # Save file
            uploaded_file.save(file_path)

            try:
                # Analyze the bill
                logger.info(f"Analyzing bill: {filename}")
                bill_analysis = analyze_bill(file_path, log_level=logging.INFO)
                results.append(bill_analysis)
            except Exception as analysis_error:
                logger.error(f"Bill analysis error for {filename}: {str(analysis_error)}")
                results.append({
                    'success': False,
                    'error': str(analysis_error),
                    'filename': filename
                })

        # Return results
        return jsonify({
            'success': len(results) > 0,
            'results': results
        }), 200

    except Exception as e:
        # Comprehensive error handling
        logger.error(f"Water tax API error: {str(e)}", exc_info=True)
        return jsonify({
            'success': False, 
            'error': str(e),
            'results': []
        }), 400

@app.route('/api/bill-chat', methods=['POST'])
def bill_chat_api():
    """
    API endpoint for bill chat using Ollama
    
    Supports:
    - Bill text analysis
    - Multilingual responses
    - Context-aware interactions
    """
    try:
        # Get request data
        data = request.get_json()
        
        # Validate input
        if not data or 'message' not in data:
            return jsonify({
                'success': False, 
                'error': 'No message provided'
            }), 400
        
        # Extract message and context
        message = data.get('message', '')
        context = data.get('context', None)
        
        # Instruction for bill analysis
        instruction = "bill information + english text and use turkish liras for pricemen and give advice for water saving tips"
        
        # Create water bill chat instance
        bill_chat = create_water_bill_chat()
        
        # Analyze bill text
        analysis = bill_chat.analyze_water_bill(message, instruction)
        
        # Return response
        return jsonify({
            'success': True,
            'response': analysis.get('response', 'No analysis available'),
            'context': analysis.get('context', [])
        })
    
    except Exception as e:
        logger.error(f"Bill chat API error: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/bill-details', methods=['POST'])
def bill_details_analysis():
    """
    Analyze bill details with comprehensive insights
    """
    try:
        # Verify CSRF token
        if not verify_csrf_token():
            return jsonify({'error': 'Invalid CSRF token'}), 403
        
        # Parse incoming request data
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No bill text provided'}), 400
        
        bill_text = data.get('bill_text', '').strip()
        
        if not bill_text:
            return jsonify({'error': 'Empty bill text'}), 400
        
        # Use bill analysis function
        from utils.advanced_bill_analyzer import process_bill_with_advanced_analysis
        
        # Analyze bill
        bill_analysis = process_bill_with_advanced_analysis(bill_text)
        
        # Create Water Expert instance
        water_expert = create_water_expert('bill_analysis')
        
        # Generate explanation
        explanation, error = water_expert.generate_bill_explanation(bill_analysis)
        
        if error:
            return jsonify({
                'error': 'Failed to generate bill explanation',
                'details': error
            }), 500
        
        return jsonify({
            'bill_details': bill_analysis,
            'explanation': explanation
        })
    
    except Exception as e:
        logger.error(f"Bill details analysis error: {e}")
        logger.error(traceback.format_exc())
        return jsonify({
            'error': 'Internal server error',
            'details': str(e)
        }), 500

@app.route('/api/chat', methods=['POST'])
def main_chat():
    try:
        data = request.get_json()
        user_message = data.get('message', '')
        
        # Validate input
        if not user_message:
            return jsonify({"error": "No message provided"}), 400
        
        # Initialize chatbot
        try:
            bot = initialize_chatbot()
        except Exception as init_error:
            logger.error(f"Chatbot initialization error: {init_error}")
            return jsonify({
                'error': 'AI service initialization failed',
                'response': 'I apologize, but the AI service is currently unavailable.'
            }), 503
        
        # Generate response
        try:
            # Verify generate_response method exists
            if not hasattr(bot, 'generate_response'):
                logger.error("Chatbot does not have generate_response method")
                return jsonify({"error": "Chatbot method not found"}), 500

            # Detailed logging before method call
            logger.info(f"Calling generate_response with context: {user_message[:200]}...")
            
            response, error = bot.generate_response(user_message)
            
            # Log generation results
            logger.info(f"Response generated: {response is not None}")
            logger.info(f"Error: {error}")

            # Handle response
            if error:
                logger.warning(f"Generation error: {error}")
                return jsonify({"error": error}), 400
            
            if not response:
                logger.warning("Empty response generated")
                response = "I apologize, but I couldn't generate a meaningful response."

            return jsonify({
                "response": response
            }), 200

        except Exception as e:
            # Comprehensive error logging
            logger.error(f"Unexpected chatbot generation error: {e}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return jsonify({"error": "Internal server error during response generation"}), 500

    except Exception as general_error:
        # Catch any other unexpected errors
        logger.error(f"Unexpected error in main chat API: {general_error}")
        return jsonify({"error": "Unexpected server error"}), 500

@app.route('/chat', methods=['POST'])
def generic_chat_api():
    """
    Generic chat endpoint that handles various chat interactions.
    
    Supports:
    - Text-based chat
    - Role-based responses
    - Fallback mechanism
    """
    try:
        data = request.get_json()
        
        # Extract necessary information
        message = data.get('message', '')
        role = data.get('role', 'general')
        
        # Validate input
        if not message:
            return jsonify({"error": "No message provided"}), 400
        
        # Determine appropriate chat handler based on role
        chat_handlers = {
            'educator': educator_chat_api,
            'farmer': farmer_chat_api,
            'general': main_chat_api
        }
        
        # Select handler, default to main chat if role not recognized
        handler = chat_handlers.get(role, main_chat_api)
        
        # Call the appropriate chat handler
        return handler()
    
    except Exception as e:
        logger.error(f"Error in generic chat API: {e}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/api/switch-model', methods=['POST'])
def switch_model_api():
    """
    API endpoint to dynamically switch the AI model
    """
    try:
        from model_inference_education import WaterConservationBot
        
        # Get the new model name from the request
        data = request.get_json()
        new_model_name = data.get('model_name')
        
        if not new_model_name:
            return jsonify({"error": "No model name provided"}), 400
        
        # Initialize the bot and attempt to switch model
        water_expert = WaterConservationBot()
        switch_success = water_expert.switch_model(new_model_name)
        
        if switch_success:
            return jsonify({
                "message": f"Successfully switched to model: {new_model_name}",
                "current_model": new_model_name
            }), 200
        else:
            return jsonify({
                "error": f"Failed to switch to model: {new_model_name}",
                "current_model": water_expert.model_name
            }), 500
    
    except Exception as e:
        logger.error(f"Error in model switching: {str(e)}")
        return jsonify({
            "error": "Unexpected error during model switch", 
            "details": str(e)
        }), 500

@app.route('/graph-techniques')
def graph_techniques():
    """
    Render the graph techniques page
    """
    try:
        return render_template('graph_techniques.html')
    except Exception as e:
        logger.error(f"Error rendering graph techniques page: {str(e)}")
        return render_template('error.html', error_message="Unable to load graph techniques page"), 500

@app.route('/generate-graph')
def generate_graph():
    return render_template('graph_generator.html')

@app.route('/generate-graph-api', methods=['POST'])
def generate_graph_api():
    data = request.get_json()
    prompt = data.get('prompt', '')
    
    result = generate_graph_from_prompt(prompt)
    return jsonify(result)

@app.route('/analyze-data', methods=['POST'])
def analyze_data():
    try:
        # Get data from request
        data = request.get_json()
        
        # Initialize Ollama Data Analyzer
        analyzer = OllamaDataAnalyzer()
        
        # Process and analyze data
        results = analyzer.process_and_analyze(
            data.get('data', {}),
            data.get('analysis_prompt'),
            data.get('chart_type', 'bar')
        )
        
        return jsonify({
            'status': 'success',
            'textual_analysis': results['textual_analysis'],
            'visualization_path': results['visualization_path']
        })
    
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

def generate_graph_from_prompt(prompt):
    """
    Generate a graph based on a user's text prompt using Ollama
    
    :param prompt: User's text description of the graph
    :return: Dictionary with graph generation results
    """
    try:
        # Initialize Ollama Data Analyzer
        analyzer = OllamaDataAnalyzer()
        
        # Generate graph details
        result = analyzer.generate_graph_from_prompt(prompt)
        
        return result
    
    except Exception as e:
        return {
            "status": "error",
            "message": f"Error in graph generation: {str(e)}"
        }

app.register_blueprint(weather_bp)

@app.route('/new-chat-endpoint', methods=['POST'])
def new_chat_endpoint():
    """
    Enhanced chat endpoint with robust error handling and simplified processing
    """
    # Minimal request logging
    app.logger.info(f"Received chat request from {request.remote_addr}")
    
    # Validate request content type
    if not request.is_json:
        app.logger.warning("Invalid request: Not JSON")
        return jsonify({"error": "Invalid request format"}), 400

    # Parse JSON data with error handling
    try:
        data = request.get_json()
    except Exception as json_error:
        app.logger.error(f"JSON parsing error: {json_error}")
        return jsonify({"error": "Invalid JSON payload"}), 400

    # Extract and validate message
    message = data.get('message', '').strip()
    file_content = data.get('file_content')
    
    if not message:
        app.logger.warning("Empty message received")
        return jsonify({"error": "Message cannot be empty"}), 400

    # Ensure chatbot is initialized
    if not hasattr(app, 'chatbot'):
        app.logger.error("Chatbot not initialized")
        return jsonify({"error": "Chatbot service unavailable"}), 500

    # Prepare context
    full_context = message
    if file_content:
        full_context += f"\nAdditional context: {file_content}"

    # Response generation with timeout protection
    try:
        # Verify generate_response method exists
        if not hasattr(app.chatbot, 'generate_response'):
            app.logger.error("Invalid chatbot configuration")
            return jsonify({"error": "Chatbot configuration error"}), 500

        # Attempt response generation with timeout
        from concurrent.futures import ThreadPoolExecutor, TimeoutError

        with ThreadPoolExecutor() as executor:
            future = executor.submit(app.chatbot.generate_response, full_context)
            
            try:
                response, error = future.result(timeout=30)  # 30-second timeout
            except TimeoutError:
                app.logger.warning("Response generation timed out")
                return jsonify({"error": "Response generation timed out"}), 504

        # Handle response variations
        if error:
            app.logger.warning(f"Generation error: {error}")
            return jsonify({"error": error}), 400
        
        if not response:
            app.logger.warning("Empty response generated")
            response = "I apologize, but I couldn't generate a meaningful response."

        return jsonify({
            "message": response,
            "source": getattr(app.chatbot, 'model_name', 'unknown')
        })

    except Exception as e:
        # Comprehensive but concise error logging
        app.logger.error(f"Unexpected chatbot error: {e}")
        return jsonify({"error": "Internal server error during response generation"}), 500

@app.route('/chat', methods=['POST'])
def chat_endpoint():
    """
    Enhanced chat endpoint with file upload support.
    
    Handles:
    - Text messages
    - Multiple file uploads
    - Context-aware interactions
    """
    try:
        # Check for message in form data
        message = request.form.get('message', '')
        context = request.form.get('context', 'general')
        
        # Process uploaded files
        uploaded_files = process_uploaded_files(request)
        
        # Prepare context for file processing
        file_context = ""
        if uploaded_files:
            file_context = "User uploaded the following files:\n"
            for file_info in uploaded_files:
                file_context += f"- {file_info['original_name']} (Type: {file_info['mime_type']}, Size: {file_info['size']} bytes)\n"
        
        # Combine message with file context
        full_message = f"{file_context}\n{message}".strip()
        
        # Create Water Expert instance
        water_expert = create_water_expert(context)
        
        # Generate response
        response, error = water_expert.generate_response(full_message)
        
        if error:
            return jsonify({"error": error}), 400
        
        return jsonify({
            "response": response,
            "files": [f['original_name'] for f in uploaded_files]
        })
    
    except Exception as e:
        logger.error(f"Chat endpoint error: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": "Internal server error"}), 500

def create_water_expert(context='general'):
    """
    Create a water expert instance based on context.
    
    Args:
        context (str): Specific context for the water expert
    
    Returns:
        An instance of the appropriate water expert
    """
    from model_inference_advanced import WaterConservationBot as WaterExpert
    from model_inference_farmers import WaterConservationBot as FarmersExpert
    from model_inference_education import WaterConservationBot as EducationExpert
    
    context_map = {
        'farmer': FarmersExpert,
        'education': EducationExpert,
        'general': WaterExpert
    }
    
    expert_class = context_map.get(context, WaterExpert)
    
    return expert_class(context)

@app.route('/weather-details')
def weather_details():
    return render_template('weather_details.html')

@app.route('/api/weather-details')
def api_weather_details():
    """
    Fetch comprehensive weather details for user's location using Open-Meteo API
    """
    try:
        # Get latitude and longitude from request
        lat = request.args.get('lat')
        lon = request.args.get('lon')

        # Log the received coordinates
        app.logger.info(f"Received coordinates - Latitude: {lat}, Longitude: {lon}")

        if not lat or not lon:
            app.logger.error("Missing latitude or longitude")
            return jsonify({
                'success': False, 
                'error': 'Missing location coordinates', 
                'details': 'Latitude and longitude are required'
            }), 400

        # Construct Open-Meteo API URL with comprehensive parameters
        url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current=temperature_2m,relative_humidity_2m,wind_speed_10m,weather_code,soil_temperature_0cm,soil_moisture_0_1cm&hourly=temperature_2m,precipitation_probability&daily=weather_code,temperature_2m_max,temperature_2m_min,uv_index_max,windspeed_10m_max,precipitation_sum&timezone=auto&forecast_days=7"

        # Log the constructed URL
        app.logger.info(f"Open-Meteo API Request URL: {url}")

        # Make the API request
        response = requests.get(url, timeout=10)
        
        # Log the full response for debugging
        app.logger.info(f"Open-Meteo API Response Status: {response.status_code}")
        app.logger.debug(f"Open-Meteo API Full Response: {response.text}")

        # Check if the request was successful
        if response.status_code != 200:
            app.logger.error(f"Open-Meteo API Error: {response.status_code} - {response.text}")
            return jsonify({
                'success': False, 
                'error': 'Unable to fetch weather data', 
                'details': f'API returned status {response.status_code}'
            }), 500

        # Parse the JSON response
        try:
            data = response.json()
        except json.JSONDecodeError as e:
            app.logger.error(f"JSON Decode Error: {e}")
            return jsonify({
                'success': False, 
                'error': 'Invalid JSON response', 
                'details': str(e)
            }), 500

        # Validate the response structure
        if not data:
            app.logger.error("Empty response received from Open-Meteo API")
            return jsonify({
                'success': False, 
                'error': 'Empty weather data', 
                'details': 'No data received from weather service'
            }), 500

        # Log the received data keys for debugging
        app.logger.info(f"Received data keys: {list(data.keys())}")
        if 'current' in data:
            app.logger.info(f"Current data keys: {list(data['current'].keys())}")
            app.logger.info(f"Current data content: {data['current']}")
        if 'daily' in data:
            app.logger.info(f"Daily data keys: {list(data['daily'].keys())}")

        # Ensure all required sections are present
        required_sections = ['current', 'daily', 'hourly']
        for section in required_sections:
            if section not in data:
                app.logger.error(f"Missing required section: {section}")
                return jsonify({
                    'success': False, 
                    'error': 'Incomplete weather data', 
                    'details': f"Missing {section} section"
                }), 500

        current = data['current']
        daily = data['daily']
        hourly = data['hourly']

        # Validate required current weather keys
        current_required_keys = [
            'temperature_2m', 'relative_humidity_2m', 'wind_speed_10m', 
            'weather_code', 'soil_temperature_0cm', 'soil_moisture_0_1cm'
        ]
        for key in current_required_keys:
            if key not in current:
                app.logger.error(f"Missing key in current weather data: {key}")
                return jsonify({
                    'success': False, 
                    'error': 'Unable to fetch weather data', 
                    'details': f"Missing current weather key: '{key}'"
                }), 500

        # Extract current weather details
        temp = current['temperature_2m']
        humidity = current['relative_humidity_2m']
        wind_speed = current['wind_speed_10m']
        soil_temp = current['soil_temperature_0cm']
        soil_moisture = current['soil_moisture_0_1cm']
        current_weather_code = current['weather_code']

        # Prepare forecast data
        forecast_data = []
        for i in range(min(7, len(daily.get('time', [])))):
            forecast_data.append({
                'date': datetime.fromisoformat(daily['time'][i]).strftime('%Y-%m-%d'),
                'temperature_max': round(daily['temperature_2m_max'][i], 1),
                'temperature_min': round(daily['temperature_2m_min'][i], 1),
                'precipitation_prob': round(hourly['precipitation_probability'][i * 24], 1),
                'uv_index': round(daily['uv_index_max'][i], 1),
                'wind_speed': round(daily['windspeed_10m_max'][i], 1),
                'icon': get_weather_icon(daily['weather_code'][i])
            })

        # Temperature trend
        temperature_trend = {
            'labels': [datetime.fromisoformat(t).strftime('%A') for t in daily.get('time', [])[:7]],
            'data': [(daily['temperature_2m_max'][i] + daily['temperature_2m_min'][i]) / 2 for i in range(min(7, len(daily.get('time', []))))]
        }
        
        # Precipitation data
        precipitation_data = {
            'labels': [datetime.fromisoformat(t).strftime('%A') for t in daily.get('time', [])[:7]],
            'data': [p for p in daily.get('precipitation_sum', [0] * 7)[:7]]
        }

        # Generate comprehensive recommendations
        recommendations = generate_comprehensive_recommendations(
            temp, 
            humidity, 
            wind_speed, 
            soil_temp, 
            soil_moisture,
            forecast_data
        )

        # Log successful data retrieval
        app.logger.info("Successfully retrieved and processed weather data")

        return jsonify({
            'success': True,
            'data': {
                'current': {
                    'temperature': round(temp, 1),
                    'humidity': round(humidity, 1),
                    'windSpeed': round(wind_speed, 1),
                    'description': get_weather_icon(current_weather_code),
                    'soilTemperature': round(soil_temp, 1),
                    'soilMoisture': round(soil_moisture, 1)
                },
                'forecast': forecast_data,
                'temperatureTrend': temperature_trend,
                'precipitationData': precipitation_data,
                'recommendations': recommendations
            }
        })

    except Exception as e:
        # Catch any unexpected errors
        app.logger.error(f"Unexpected error in weather details: {str(e)}")
        app.logger.error(traceback.format_exc())
        
        return jsonify({
            'success': False, 
            'error': 'Unexpected server error', 
            'details': str(e)
        }), 500

def get_weather_icon(weather_code):
    """
    Map Open-Meteo weather codes to Font Awesome icons
    """
    weather_icons = {
        0: 'fas fa-sun',      # Clear sky
        1: 'fas fa-cloud-sun', # Mainly clear
        2: 'fas fa-cloud-sun', # Partly cloudy
        3: 'fas fa-cloud',     # Overcast
        45: 'fas fa-smog',     # Foggy
        48: 'fas fa-smog',     # Depositing rime fog
        51: 'fas fa-cloud-rain', # Light drizzle
        53: 'fas fa-cloud-rain', # Moderate drizzle
        55: 'fas fa-cloud-rain', # Dense drizzle
        61: 'fas fa-cloud-rain', # Slight rain
        63: 'fas fa-cloud-rain', # Moderate rain
        65: 'fas fa-cloud-showers-heavy', # Heavy rain
        71: 'fas fa-snowflake', # Slight snow fall
        73: 'fas fa-snowflake', # Moderate snow fall
        75: 'fas fa-snowflake', # Heavy snow fall
        77: 'fas fa-snowflake', # Snow grains
        80: 'fas fa-cloud-rain', # Slight rain showers
        81: 'fas fa-cloud-rain', # Moderate rain showers
        82: 'fas fa-cloud-showers-heavy', # Violent rain showers
        85: 'fas fa-snowflake', # Slight snow showers
        86: 'fas fa-snowflake', # Heavy snow showers
        95: 'fas fa-bolt',     # Thunderstorm
        96: 'fas fa-bolt',     # Thunderstorm with light hail
        99: 'fas fa-bolt'      # Thunderstorm with heavy hail
    }
    return weather_icons.get(weather_code, 'fas fa-cloud')

def generate_comprehensive_recommendations(temp, humidity, wind_speed, soil_temp, soil_moisture, forecast):
    """
    Generate comprehensive agricultural recommendations
    """
    recommendations = {
        'irrigation': 'Standard irrigation recommended',
        'planting': 'Monitor local crop conditions',
        'harvest': 'No immediate harvest alerts'
    }

    # Irrigation recommendations
    if temp > 30 and humidity < 40:
        recommendations['irrigation'] = 'Increase irrigation due to high temperature and low humidity'
    elif soil_moisture < 20:
        recommendations['irrigation'] = 'Critical irrigation needed - soil moisture is very low'
    elif temp < 15:
        recommendations['irrigation'] = 'Reduce irrigation, minimal evaporation'

    # Planting recommendations
    if 20 <= temp <= 30 and soil_temp > 15:
        recommendations['planting'] = 'Optimal conditions for planting most crops'
    elif wind_speed > 20:
        recommendations['planting'] = 'Avoid planting due to high winds'

    # Harvest recommendations
    if forecast and len(forecast) > 0:
        next_week_avg_temp = sum(day['temperature_avg'] for day in forecast) / len(forecast)
        if next_week_avg_temp > 25:
            recommendations['harvest'] = 'Consider early morning harvesting to avoid heat stress'

    return recommendations

@app.route('/process_file_education', methods=['POST'])
def process_file_education():
    """
    Process file uploads for the educator chat interface
    Supports various file types and uses Ollama for analysis
    """
    # Check if file is present
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    message = request.form.get('message', '')
    
    # Validate file
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    # Secure filename and generate unique path
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"education_{uuid.uuid4()}_{filename}")
    
    # Save file
    file.save(file_path)
    
    try:
        # Read file content based on type
        file_type = get_file_type(filename)
        file_content = read_file_content(file_path)
        
        # Combine file content with user message
        full_prompt = f"""Analyze the following {file_type} file in the context of water conservation education:

File Content:
{file_content}

User Message: {message}

Please provide an educational analysis focusing on water conservation insights."""
        
        # Use education-specific chatbot
        from model_inference_education import WaterConservationBot
        bot = WaterConservationBot(model_name="water-expert-advanced")
        
        # Generate response
        response, error = bot.generate_response(full_prompt)
        
        if error:
            return jsonify({'error': error}), 500
        
        return jsonify({
            'response': response,
            'file_name': filename,
            'file_type': file_type
        })
    
    except Exception as e:
        logger.error(f"File processing error: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up uploaded file
        if os.path.exists(file_path):
            os.remove(file_path)

@app.route('/process_file_farmer', methods=['POST'])
def process_file_farmer():
    """Process file uploads for farmer chat"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        # Check if filename is empty
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        
        # Check file type and size
        allowed_extensions = {'.pdf', '.doc', '.docx', '.txt'}
        max_file_size = 50 * 1024 * 1024  # 50MB
        
        # Get file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            return jsonify({'error': f'File type {file_ext} not allowed'}), 400
        
        if file.content_length > max_file_size:
            return jsonify({'error': 'File too large'}), 400
        
        # Get optional message
        message = request.form.get('message', '')
        
        # Save file temporarily
        temp_dir = os.path.join(app.instance_path, 'temp_uploads')
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(temp_path)
        
        # Process file with Ollama
        from model_inference_farmers import process_farmer_file
        
        # Generate response
        response_text = process_farmer_file(temp_path, message)
        
        # Clean up temporary file
        os.remove(temp_path)
        
        return jsonify({
            'response': response_text
        })
    
    except Exception as e:
        # Log the error
        logger.error(f"Error processing farmer file: {e}", exc_info=True)
        return jsonify({
            'error': 'An error occurred while processing the file',
            'details': str(e)
        }), 500

@app.route('/get-weather')
def get_weather_route():
    """
    Route wrapper for get_weather function to handle location and weather retrieval.
    """
    try:
        # Call the original get_weather function
        result = get_weather()
        
        # Log the result for debugging
        app.logger.info(f"Get Weather Result: {result}")
        
        return jsonify(result)
    
    except Exception as e:
        # Log any unexpected errors
        app.logger.error(f"Unexpected error in get-weather route: {e}")
        return jsonify({
            'error': 'Failed to retrieve weather information',
            'details': str(e)
        }), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
    conn = sqlite3.connect('feedback.db')
    c = conn.cursor()
    
    # Create feedback table
    c.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            email TEXT,
            rating INTEGER,
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create admin table
    c.execute('''
        CREATE TABLE IF NOT EXISTS admin (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT
        )
    ''')
    
    # Add default admin if not exists
    c.execute('SELECT * FROM admin WHERE username = ?', ('admin',))
    if not c.fetchone():
        c.execute('INSERT INTO admin (username, password) VALUES (?, ?)',
                 ('admin', generate_password_hash('admin123')))
    
    conn.commit()
    conn.close()

# Initialize database on startup
init_db()

# Load translations
with open('translations.json', 'r', encoding='utf-8') as f:
    translations = json.load(f)

def get_text(key):
    """Get text in English"""
    try:
        return translations['en'].get(key, key)
    except:
        return key

# Register get_text as a template filter
app.jinja_env.filters['translate'] = get_text

@app.context_processor
def utility_processor():
    return dict(get_text=get_text)

# Sanitize input
def sanitize_input(text):
    return bleach.clean(text, strip=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_type(filename):
    """Get file type based on extension"""
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    return mimetypes.guess_type(filename)[0] or 'application/octet-stream'

def read_file_content(file_path):
    """
    Read file content based on file type with robust error handling
    Supports multiple file types: txt, pdf, docx, images
    
    Args:
        file_path (str): Path to the uploaded file
    
    Returns:
        str: Extracted text content from the file
    """
    try:
        # Determine file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        # Text files
        if ext in ['.txt', '.csv']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        # PDF files
        elif ext == '.pdf':
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text

        # Word documents
        elif ext in ['.doc', '.docx']:
            import docx
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs if para.text])

        # Image files (OCR)
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            import pytesseract
            from PIL import Image
            
            # Ensure Tesseract is installed
            try:
                return pytesseract.image_to_string(Image.open(file_path))
            except Exception as ocr_error:
                logger.warning(f"OCR failed for image: {ocr_error}")
                return "Could not extract text from image"

        else:
            logger.warning(f"Unsupported file type: {ext}")
            return f"Unsupported file type: {ext}"

    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        return f"Error reading file: {str(e)}"

def generate_csrf_token():
    if 'csrf_token' not in session:
        session['csrf_token'] = secrets.token_hex(32)
    return session['csrf_token']

app.jinja_env.globals['csrf_token'] = generate_csrf_token

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in to access this page', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def verify_csrf_token():
    # Skip CSRF check for local network requests
    if request.remote_addr.startswith('192.168.') or request.remote_addr == '127.0.0.1':
        return True
        
    # Check form data first
    token = request.form.get('csrf_token')
    if not token and request.is_json:
        # Then check JSON data
        token = request.json.get('csrf_token')
    # Finally check headers
    if not token:
        token = request.headers.get('X-CSRF-Token')
    return token and token == session.get('csrf_token')

@app.before_request
def before_request():
    # Skip session check for static files and public routes
    if request.path.startswith('/static') or request.endpoint in ['login', 'index', 'logout']:
        return

    # Only check admin routes
    if request.endpoint and request.endpoint.startswith('admin'):
        if 'user_id' not in session:
            flash('Please log in to access the admin panel.', 'error')
            return redirect(url_for('login'))
        
        # Check session timeout (20 minutes)
        if 'last_activity' in session:
            last_activity = datetime.fromisoformat(session['last_activity'])
            if (datetime.now() - last_activity).total_seconds() > 1200:  # 20 minutes
                session.clear()
                flash('Your session has expired. Please log in again.', 'error')
                return redirect(url_for('login'))
        
        # Update last activity time
        session['last_activity'] = datetime.now().isoformat()

    # Require CSRF token for all POST requests
    if request.method == 'POST':
        if not verify_csrf_token():
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'error': 'Invalid CSRF token'}), 400
            flash('Invalid CSRF token. Please try again.', 'error')
            return redirect(url_for('index'))